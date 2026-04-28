"""生成《信道生成与采集-物理层建模梳理》Word 文档。

运行: python docs/gen_channel_modeling_doc.py
输出: docs/信道生成与采集-物理层建模梳理.docx
"""

from __future__ import annotations

import io
import math
import os
import textwrap
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# 全局配置
# ---------------------------------------------------------------------------
OUT_DIR = Path(__file__).resolve().parent
IMG_DIR = OUT_DIR / "_img_tmp"
IMG_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "信道生成与采集-物理层建模梳理.docx"

plt.rcParams.update({
    "font.family": "Microsoft YaHei",
    "font.size": 9,
    "axes.unicode_minus": False,
    "figure.dpi": 180,
})

# ---------------------------------------------------------------------------
# 插图 1: 信道生成因果链路图
# ---------------------------------------------------------------------------
def fig1_causal_chain() -> Path:
    fig, ax = plt.subplots(figsize=(10, 7.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis("off")

    box_style = dict(boxstyle="round,pad=0.35", facecolor="#E3F2FD", edgecolor="#1565C0", linewidth=1.2)
    box_style2 = dict(boxstyle="round,pad=0.35", facecolor="#FFF3E0", edgecolor="#E65100", linewidth=1.2)
    box_style3 = dict(boxstyle="round,pad=0.35", facecolor="#E8F5E9", edgecolor="#2E7D32", linewidth=1.2)
    arrow_kw = dict(arrowstyle="->,head_width=0.15,head_length=0.1", color="#333", lw=1.2)
    note_kw = dict(fontsize=6.5, color="#666", fontstyle="italic")

    boxes = {
        "net":   (2.0, 7.2, "组网建模\n站点/扇区/UE位置", box_style),
        "mob":   (7.0, 7.2, "移动性建模\nstatic/linear/RW/RWP", box_style),
        "pl":    (2.0, 5.8, "大尺度衰落\n路径损耗 + 阴影衰落", box_style2),
        "serv":  (5.5, 5.8, "服务小区选择\nargmax(Rx功率)", box_style2),
        "dopp":  (8.2, 5.8, "多普勒频移\nfd=v·fc/c", box_style2),
        "tdl":   (5.0, 4.2, "TDL信道模型\nh [K,T,RB,Tx,Rx]", box_style3),
        "tdd":   (1.2, 3.0, "TDD时隙配置\nDDDSU等", box_style),
        "pilot": (4.0, 2.4, "导频生成\nSRS/CSI-RS", box_style),
        "intf":  (7.5, 2.4, "干扰叠加\nY=H·X+ΣHi·Xi+n", box_style2),
        "est":   (5.0, 1.0, "信道估计\nLS → MMSE → 插值", box_style3),
        "ssb":   (8.5, 1.0, "SSB测量\nRSRP/RSRQ", box_style),
        "out":   (5.0, 0.0, "ChannelSample 输出", box_style3),
    }

    for key, (x, y, txt, bs) in boxes.items():
        ax.text(x, y, txt, ha="center", va="center", fontsize=8,
                bbox=bs, fontweight="bold" if key == "out" else "normal")

    arrows = [
        ("net", "pl", "d_3d, scenario"),
        ("net", "serv", "K个小区"),
        ("mob", "dopp", "速度v"),
        ("pl", "serv", "Rx功率"),
        ("pl", "tdl", "功率缩放"),
        ("serv", "tdl", "服务/干扰划分"),
        ("dopp", "tdl", "fd"),
        ("tdl", "pilot", ""),
        ("tdd", "pilot", "可用时隙"),
        ("pilot", "intf", "X_RS"),
        ("tdl", "intf", "h_intf"),
        ("intf", "est", "Y"),
        ("pilot", "est", "导频位置"),
        ("est", "out", "h_est"),
        ("tdl", "ssb", "多小区h"),
        ("ssb", "out", "RSRP"),
    ]

    for src, dst, label in arrows:
        sx, sy = boxes[src][0], boxes[src][1]
        dx, dy = boxes[dst][0], boxes[dst][1]
        mid_x = (sx + dx) / 2
        mid_y = (sy + dy) / 2
        ax.annotate("", xy=(dx, dy + 0.25), xytext=(sx, sy - 0.25), arrowprops=arrow_kw)
        if label:
            ax.text(mid_x + 0.05, mid_y + 0.1, label, **note_kw)

    ax.set_title("图1  信道生成因果链路图 —— 模块间数据依赖与参数传递", fontsize=11, fontweight="bold", pad=10)
    path = IMG_DIR / "fig1_causal_chain.png"
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


# ---------------------------------------------------------------------------
# 插图 2: 七站三扇区组网图
# ---------------------------------------------------------------------------
def fig2_hex_topology() -> Path:
    fig, ax = plt.subplots(figsize=(7, 6.5))
    isd = 500
    hex_dirs = [(1, 0), (0.5, math.sqrt(3)/2), (-0.5, math.sqrt(3)/2),
                (-1, 0), (-0.5, -math.sqrt(3)/2), (0.5, -math.sqrt(3)/2)]

    sites = [(0, 0)]
    for d in hex_dirs:
        sites.append((d[0] * isd, d[1] * isd))

    colors_sector = ["#E53935", "#43A047", "#1E88E5"]
    sector_angles_deg = [90, 210, 330]

    for sx, sy in sites:
        hex_r = isd / math.sqrt(3)
        hex_pts = [(sx + hex_r * math.cos(math.radians(60 * i + 30)),
                     sy + hex_r * math.sin(math.radians(60 * i + 30))) for i in range(6)]
        hex_poly = plt.Polygon(hex_pts, fill=False, edgecolor="#999", lw=0.8, ls="--")
        ax.add_patch(hex_poly)
        ax.plot(sx, sy, "k^", ms=8, zorder=5)
        for si, ang in enumerate(sector_angles_deg):
            dx = 120 * math.cos(math.radians(ang))
            dy = 120 * math.sin(math.radians(ang))
            ax.annotate("", xy=(sx + dx, sy + dy), xytext=(sx, sy),
                        arrowprops=dict(arrowstyle="-|>", color=colors_sector[si], lw=1.5))

    rng = np.random.default_rng(42)
    n_ue = 40
    ue_x = rng.uniform(-700, 700, n_ue)
    ue_y = rng.uniform(-600, 600, n_ue)
    ax.scatter(ue_x, ue_y, c="#FF9800", s=25, zorder=4, edgecolors="k", linewidths=0.4, label="UE")

    for i, (sx, sy) in enumerate(sites):
        ax.text(sx + 20, sy + 25, f"Site {i}", fontsize=6.5, color="#333")

    ax.set_xlim(-850, 850)
    ax.set_ylim(-750, 750)
    ax.set_aspect("equal")
    ax.set_xlabel("x (m)")
    ax.set_ylabel("y (m)")
    ax.set_title("图2  七站三扇区六边形组网 + UE分布", fontsize=11, fontweight="bold")
    ax.legend(loc="upper right", fontsize=8)
    ax.grid(True, alpha=0.2)
    path = IMG_DIR / "fig2_hex_topology.png"
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


# ---------------------------------------------------------------------------
# 插图 3: 移动性模型轨迹
# ---------------------------------------------------------------------------
def fig3_mobility() -> Path:
    fig, axes = plt.subplots(2, 2, figsize=(8, 7))
    rng = np.random.default_rng(7)
    n_steps = 80
    dt = 0.5

    # static
    ax = axes[0, 0]
    ax.plot(100, 100, "ro", ms=10)
    ax.set_title("(a) Static", fontweight="bold")
    ax.set_xlim(0, 200); ax.set_ylim(0, 200)
    ax.text(110, 90, "UE 固定不动\nfd = 0", fontsize=7.5)
    ax.set_aspect("equal")
    ax.grid(True, alpha=0.3)

    # linear
    ax = axes[0, 1]
    speed = 3 / 3.6
    angle = math.radians(30)
    t = np.arange(n_steps) * dt
    lx = 20 + speed * np.cos(angle) * t
    ly = 20 + speed * np.sin(angle) * t
    ax.plot(lx, ly, "b-", lw=1.5, alpha=0.8)
    ax.plot(lx[0], ly[0], "go", ms=7, label="起点")
    ax.plot(lx[-1], ly[-1], "rs", ms=7, label="终点")
    ax.set_title("(b) Linear", fontweight="bold")
    ax.legend(fontsize=6.5)
    ax.set_aspect("equal")
    ax.grid(True, alpha=0.3)

    # random walk
    ax = axes[1, 0]
    wx, wy = [100.0], [100.0]
    for _ in range(n_steps):
        a = rng.uniform(0, 2 * math.pi)
        s = speed * dt
        wx.append(wx[-1] + s * math.cos(a))
        wy.append(wy[-1] + s * math.sin(a))
    ax.plot(wx, wy, "m-", lw=1, alpha=0.7)
    ax.plot(wx[0], wy[0], "go", ms=7)
    ax.plot(wx[-1], wy[-1], "rs", ms=7)
    ax.set_title("(c) Random Walk", fontweight="bold")
    ax.set_aspect("equal")
    ax.grid(True, alpha=0.3)

    # random waypoint
    ax = axes[1, 1]
    wpx, wpy = [50.0], [50.0]
    while len(wpx) < n_steps + 1:
        tx, ty = rng.uniform(0, 200), rng.uniform(0, 200)
        seg = max(1, int(math.hypot(tx - wpx[-1], ty - wpy[-1]) / (speed * dt)))
        for i in range(1, seg + 1):
            frac = i / seg
            wpx.append(wpx[-1 * seg if len(wpx) > seg else 0] * 0 + wpx[-1] + (tx - wpx[-1]) * (1/seg))
            wpy.append(wpy[-1] + (ty - wpy[-1]) * (1/seg))
            if len(wpx) > n_steps:
                break
    wpx2, wpy2 = [50.0], [50.0]
    targets = [(rng.uniform(20, 180), rng.uniform(20, 180)) for _ in range(6)]
    for tx, ty in targets:
        dist = math.hypot(tx - wpx2[-1], ty - wpy2[-1])
        steps = max(1, int(dist / (speed * dt)))
        for i in range(1, steps + 1):
            wpx2.append(wpx2[-1] + (tx - wpx2[-1]) / (steps - i + 1))
            wpy2.append(wpy2[-1] + (ty - wpy2[-1]) / (steps - i + 1))
    ax.plot(wpx2, wpy2, "c-", lw=1, alpha=0.7)
    for tx, ty in targets:
        ax.plot(tx, ty, "kx", ms=5, alpha=0.5)
    ax.plot(wpx2[0], wpy2[0], "go", ms=7)
    if len(wpx2) > 1:
        ax.plot(wpx2[-1], wpy2[-1], "rs", ms=7)
    ax.set_title("(d) Random Waypoint", fontweight="bold")
    ax.set_aspect("equal")
    ax.grid(True, alpha=0.3)

    fig.suptitle("图3  四种UE移动性模型轨迹对比", fontsize=11, fontweight="bold", y=1.01)
    path = IMG_DIR / "fig3_mobility.png"
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


# ---------------------------------------------------------------------------
# 插图 4: TDL 功率延迟分布
# ---------------------------------------------------------------------------
def fig4_tdl_pdp() -> Path:
    # TDL-C profile (approx 3GPP 38.901 Table 7.5-4)
    delays_ns = np.array([0, 65, 70, 190, 195, 200, 240, 325, 520, 590, 610,
                           730, 800, 960, 1020, 1100, 1210, 1420, 1520, 1640,
                           1725, 1800, 2170])
    powers_dB = np.array([-4.4, -1.2, -3.5, -5.2, -2.8, 0.0, -2.2, -3.9, -7.4,
                           -7.1, -10.7, -11.1, -5.1, -6.8, -8.7, -13.2, -13.9,
                           -13.9, -15.8, -17.1, -16.0, -15.7, -21.6])
    tau_rms_ns = 363

    fig, ax = plt.subplots(figsize=(8, 4))
    markerline, stemlines, baseline = ax.stem(delays_ns, powers_dB, linefmt="-", markerfmt="o", basefmt=" ")
    plt.setp(stemlines, color="#1565C0", linewidth=1.5)
    plt.setp(markerline, color="#1565C0", markersize=5)

    ax.axhline(0, color="#999", lw=0.5, ls="--")
    ax.set_xlabel("时延 τ (ns)")
    ax.set_ylabel("归一化功率 (dB)")
    ax.set_title(f"图4  TDL-C 功率延迟分布 (NLOS, τ_rms ≈ {tau_rms_ns} ns)", fontsize=11, fontweight="bold")
    ax.text(1500, -2, f"23个多径抽头\nRMS延迟扩展 ≈ {tau_rms_ns} ns", fontsize=8,
            bbox=dict(facecolor="lightyellow", edgecolor="#ccc", alpha=0.9))
    ax.grid(True, alpha=0.2)
    path = IMG_DIR / "fig4_tdl_pdp.png"
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


# ---------------------------------------------------------------------------
# 插图 5: TDD 时隙配置色块图
# ---------------------------------------------------------------------------
def fig5_tdd_pattern() -> Path:
    patterns = {
        "DDDSU":      "DDDSU",
        "DDSUU":      "DDSUU",
        "DDDDDDDSUU": "DDDDDDDSUU",
        "DSUUD":      "DSUUD",
    }
    sym_per_slot = 14
    color_map = {"D": "#42A5F5", "U": "#66BB6A", "G": "#FFA726"}
    special_split = (10, 2, 2)

    fig, axes = plt.subplots(len(patterns), 1, figsize=(10, 4.5), sharex=False)
    for idx, (name, slots) in enumerate(patterns.items()):
        ax = axes[idx]
        x_offset = 0
        for si, s in enumerate(slots):
            if s == "D":
                syms = ["D"] * sym_per_slot
            elif s == "U":
                syms = ["U"] * sym_per_slot
            else:
                syms = ["D"] * special_split[0] + ["G"] * special_split[1] + ["U"] * special_split[2]
            for sym_i, sym_type in enumerate(syms):
                rect = plt.Rectangle((x_offset + sym_i, 0), 1, 1, facecolor=color_map[sym_type],
                                     edgecolor="white", linewidth=0.3)
                ax.add_patch(rect)
            ax.text(x_offset + len(syms) / 2, -0.35, f"Slot {si}", ha="center", fontsize=5.5, color="#555")
            x_offset += len(syms)
        ax.set_xlim(0, x_offset)
        ax.set_ylim(-0.6, 1.2)
        ax.set_yticks([])
        ax.set_xticks([])
        ax.set_ylabel(name, fontsize=8, fontweight="bold", rotation=0, labelpad=65, va="center")
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_visible(False)
        ax.spines["bottom"].set_visible(False)

    legend_patches = [mpatches.Patch(color=c, label=l) for l, c in
                      [("DL", "#42A5F5"), ("UL", "#66BB6A"), ("Guard", "#FFA726")]]
    axes[0].legend(handles=legend_patches, loc="upper right", fontsize=7, ncol=3, framealpha=0.8)
    fig.suptitle("图5  TDD 时隙配置 (符号级 DL/UL/Guard 分布)", fontsize=11, fontweight="bold", y=1.0)
    path = IMG_DIR / "fig5_tdd_pattern.png"
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


# ---------------------------------------------------------------------------
# 插图 6: SRS 跳频示意图
# ---------------------------------------------------------------------------
def fig6_srs_hopping() -> Path:
    fig, ax = plt.subplots(figsize=(9, 5))
    total_rb = 273
    m_srs = 68  # each hop covers ~68 RBs (273/4)
    n_hops = 4
    k_tc = 2
    comb_label = f"K_TC={k_tc}"

    rng = np.random.default_rng(12)
    hop_colors = ["#1565C0", "#E53935", "#43A047", "#FF8F00"]
    slot_labels = []

    for hop in range(n_hops):
        rb_start = hop * m_srs
        rb_end = min(rb_start + m_srs, total_rb)
        slot_idx = hop * 2
        for slot_off in range(2):
            s = slot_idx + slot_off
            rect = plt.Rectangle((s, rb_start), 1, rb_end - rb_start,
                                  facecolor=hop_colors[hop], alpha=0.6, edgecolor="k", lw=0.5)
            ax.add_patch(rect)
            slot_labels.append(f"Slot {s}")

    ax.set_xlim(-0.5, n_hops * 2 + 0.5)
    ax.set_ylim(0, total_rb)
    ax.set_xlabel("时隙 (Slot)")
    ax.set_ylabel("RB 索引")
    ax.set_xticks(np.arange(n_hops * 2) + 0.5)
    ax.set_xticklabels([f"Slot {i}" for i in range(n_hops * 2)], fontsize=7)

    for hop in range(n_hops):
        rb_start = hop * m_srs
        rb_mid = rb_start + m_srs // 2
        ax.text(hop * 2 + 1, rb_mid, f"Hop {hop}\nRB {rb_start}-{min(rb_start+m_srs, total_rb)-1}",
                ha="center", va="center", fontsize=7, fontweight="bold", color="white")

    ax.axhline(total_rb, color="#999", lw=0.5, ls="--")
    ax.text(n_hops * 2 + 0.3, total_rb - 10, f"总RB={total_rb}", fontsize=7, color="#666")

    ax.annotate("4次跳频覆盖全频带", xy=(3.5, total_rb * 0.5),
                fontsize=9, color="#333", fontweight="bold",
                bbox=dict(facecolor="lightyellow", edgecolor="#ccc"))

    ax.set_title("图6  SRS 频率跳频示意 (带宽树 4级, 100MHz@30kHz)", fontsize=11, fontweight="bold")
    ax.grid(True, alpha=0.15, axis="y")
    path = IMG_DIR / "fig6_srs_hopping.png"
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


# ---------------------------------------------------------------------------
# 插图 7: 信道估计流程图
# ---------------------------------------------------------------------------
def fig7_channel_est() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)
    ax.axis("off")

    box_kw = dict(boxstyle="round,pad=0.3", linewidth=1.2)
    bk1 = dict(**box_kw, facecolor="#E3F2FD", edgecolor="#1565C0")
    bk2 = dict(**box_kw, facecolor="#FFF3E0", edgecolor="#E65100")
    bk3 = dict(**box_kw, facecolor="#E8F5E9", edgecolor="#2E7D32")
    arr = dict(arrowstyle="->,head_width=0.12", color="#333", lw=1.2)

    nodes = [
        (1.0, 3.0, "接收信号 Y\n= H·X + 干扰 + 噪声", bk1),
        (3.5, 3.0, "导频提取\nY_RS, X_RS", bk1),
        (5.8, 3.0, "LS估计\nĤ = Y/X", bk2),
        (8.2, 3.0, "MMSE细化\nĤ_MMSE = R(R+σ²I)⁻¹Ĥ_LS", bk2),
        (5.8, 1.2, "2D插值\n频率→时间", bk2),
        (8.5, 1.2, "全频带信道\nh_est [T,RB,Tx,Rx]", bk3),
    ]

    for x, y, txt, bk in nodes:
        ax.text(x, y, txt, ha="center", va="center", fontsize=7.5, bbox=bk)

    conns = [(1.0, 3.0, 3.5, 3.0), (3.5, 3.0, 5.8, 3.0), (5.8, 3.0, 8.2, 3.0),
             (8.2, 2.7, 5.8, 1.5), (5.8, 1.2, 8.5, 1.2)]
    for sx, sy, dx, dy in conns:
        off_x = 0.6 if dx > sx else (-0.6 if dx < sx else 0)
        off_y = 0 if abs(dx - sx) > 0.1 else (-0.3 if dy < sy else 0.3)
        ax.annotate("", xy=(dx - off_x * 0.5, dy), xytext=(sx + off_x * 0.5, sy), arrowprops=arr)

    note_kw = dict(fontsize=6, color="#888", fontstyle="italic")
    ax.text(4.6, 3.55, "导频位置\n(第6章)", **note_kw)
    ax.text(7.0, 3.55, "SNR(第4章)\nτ_rms(第5章)", **note_kw)
    ax.text(4.5, 0.65, "导频间距→\n全RB/全符号", **note_kw)

    ax.set_title("图7  信道估计流程 (LS → MMSE → 2D插值)", fontsize=11, fontweight="bold", pad=10)
    path = IMG_DIR / "fig7_channel_est.png"
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


# ===========================================================================
# Word 文档组装
# ===========================================================================

def _add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)

def _add_para(doc: Document, text: str, bold: bool = False, indent_cm: float = 0):
    p = doc.add_paragraph()
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.name = "Microsoft YaHei"
    return p

def _add_image(doc: Document, img_path: Path, width_inches: float = 6.0):
    doc.add_picture(str(img_path), width=Inches(width_inches))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    return table


def build_document(fig_paths: dict[str, Path]):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ===== 封面 =====
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("MSG 平台\n信道生成与采集\n物理层建模梳理")
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    doc.add_paragraph()
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run("—— 以因果链串联各建模环节，讲清模块间的依赖与影响 ——")
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # ===== 第 1 章 =====
    _add_heading(doc, "第 1 章  信道生成全链路总览")
    _add_para(doc, (
        "信道采集不是各个独立模块的简单拼装，而是一条从物理场景到数据样本的因果链。"
        "组网决定了站间距和小区数量，站间距决定了路径损耗，路径损耗决定了 SNR 和服务小区选择；"
        "UE 的移动性驱动多普勒频移，多普勒频移决定了 TDL 信道的时变速率；"
        "TDD 时隙配置约束了导频只能放在特定符号，导频的频率跳频需要在信道相干时间内覆盖全频带；"
        "邻区干扰由组网拓扑和邻区导频共同决定，最终影响信道估计的精度。"
        "\n\n下图展示了这条因果链中各环节的数据依赖和参数传递关系。"
    ))
    _add_image(doc, fig_paths["fig1"], 6.2)
    _add_para(doc, (
        "关键源文件一览：组网建模 → internal_sim.py (_build_sites, _place_ues)；"
        "移动性 → _mobility.py；大尺度衰落 → internal_sim.py (_pathloss_*)；"
        "TDL → internal_sim.py (_generate_tdl_channel)；TDD → phy_sim/tdd_config.py；"
        "导频 → ref_signals/srs.py, csi_rs.py, ssb.py；"
        "干扰 → _interference_estimation.py；信道估计 → channel_est/pipeline.py, ls.py, mmse.py, interpolate.py。"
    ))
    doc.add_page_break()

    # ===== 第 2 章 =====
    _add_heading(doc, "第 2 章  组网与场景建模 —— 一切的起点")
    _add_para(doc, (
        "组网是信道生成的第一步，也是后续所有环节的输入源头。"
        "站点位置决定了每个 UE 到各基站的三维距离 d_3d，这直接送入路径损耗计算；"
        "扇区数量决定了小区总数 K，进而决定了干扰小区的集合大小；"
        "UE 分布模式影响 UE 与服务基站的距离统计特性，以及 LOS/NLOS 的概率分布。"
    ))

    _add_heading(doc, "2.1  六边形网格拓扑", level=2)
    _add_para(doc, (
        "系统采用经典的六边形蜂窝组网。中心 Ring 0 包含 1 个站点，Ring r 新增 6r 个站点，"
        "累积数量为 1 + 3r(r+1)。默认 7 站（1 ring）时恰好形成中心 + 6 邻站的经典拓扑。"
        "站间距 (ISD) 决定了六边形边长 = ISD/√3，典型值 500m (UMa) 或 200m (UMi)。"
    ))
    _add_para(doc, (
        "扇区化将每个站点分为 1 或 3 个扇区。三扇区时方位角分别为 0°、120°、240°，"
        "每个扇区作为独立小区，7 站 3 扇区 = 21 个小区。"
        "每个小区有独立的 PCI (物理小区ID)，用于导频序列的种子初始化。"
    ))

    _add_heading(doc, "2.2  UE 放置与服务小区选择", level=2)
    _add_para(doc, (
        "UE 放置支持三种模式：\n"
        "  · uniform：在网络覆盖范围内均匀随机撒点\n"
        "  · clustered：UE 集中在若干热点区域\n"
        "  · hotspot：在特定小区附近高密度放置\n\n"
        "服务小区选择采用最大接收功率准则：对每个 UE 计算其到所有 K 个小区的接收功率 "
        "(发射功率 - 路径损耗)，取最大值对应的小区作为服务小区，其余 K-1 个为潜在干扰小区。"
    ))
    _add_image(doc, fig_paths["fig2"], 5.5)

    _add_heading(doc, "2.3  组网参数对下游的影响", level=2)
    _add_table(doc,
        ["组网参数", "直接影响", "影响的下游环节"],
        [
            ["站间距 isd_m", "基站-UE 距离 d_3d", "路径损耗(第4章) → SNR → 信道估计精度(第8章)"],
            ["扇区数 sectors_per_site", "小区总数 K", "干扰小区数 K-1 → 干扰强度(第7章) → SIR"],
            ["UE 分布 ue_distribution", "UE 与各基站距离统计", "LOS/NLOS 概率(第4章)、服务小区选择"],
            ["场景 scenario", "路损公式 + TDL profile", "大尺度衰落(第4章) + 小尺度衰落(第5章)"],
            ["天线数 num_bs_tx_ant", "空间自由度", "天线相关性(第5章) + SSB波束数(第9章)"],
            ["发射功率 tx_power_dbm", "接收功率绝对值", "SNR + 服务小区选择"],
        ]
    )
    doc.add_page_break()

    # ===== 第 3 章 =====
    _add_heading(doc, "第 3 章  移动性建模 —— 驱动时变特性")
    _add_para(doc, (
        "UE 的运动不仅改变其地理位置（从而影响路径损耗的逐快照更新），更关键的是产生多普勒频移 fd，"
        "这是 TDL 信道时变特性的直接驱动力。同时，UE 移动速度还约束了 SRS 跳频的设计——"
        "跳频周期必须在信道相干时间 Tc 内完成全频带覆盖，否则不同频段的信道估计对应不同时刻的信道状态。"
    ))

    _add_heading(doc, "3.1  四种移动模型", level=2)
    _add_para(doc, (
        "· static：UE 固定不动，fd = 0，适用于室内固定终端场景。\n\n"
        "· linear：恒速直线运动，速度和方向角由配置指定。产生恒定的多普勒频移 "
        "fd = v·fc/c·cos(θ)，其中 θ 为运动方向与到达角的夹角。\n\n"
        "· random_walk：每个时间步随机选择新的运动方向，步长由速度×时间步长决定。"
        "多普勒频移随方向变化而波动。\n\n"
        "· random_waypoint：随机选取目标点，UE 沿直线匀速移向目标，到达后可停留一段时间再选新目标。"
        "这是最贴近真实行人运动的模型。"
    ))
    _add_image(doc, fig_paths["fig3"], 5.8)

    _add_heading(doc, "3.2  移动性与其他模块的耦合", level=2)
    _add_table(doc,
        ["耦合关系", "机制", "实际影响"],
        [
            ["移动性 → 多普勒 (第5章)",
             "fd = v·fc/c·cos(θ)，直接送入 TDL 的 Jakes 和弦和模型",
             "fd 越大 → 信道时变越快 → 相邻时隙间信道差异越大"],
            ["移动性 → 路损时变 (第4章)",
             "UE 位置 x(t) 变化 → d_3d(t) 变化 → 路损逐快照更新",
             "UE 远离基站时 SNR 下降，靠近时 SNR 上升"],
            ["移动性 → LSP 空间相关",
             "轨迹上相邻位置的大尺度参数(延迟扩展、角度扩展等)有空间连续性",
             "保证移动过程中信道统计特性平滑变化，而非跳变"],
            ["移动性 → SRS 跳频 (第6章)",
             "信道相干时间 Tc ∝ 1/fd；跳频周期 T_SRS 需 < Tc",
             "高速 UE 需更快完成全频带覆盖，否则频段间信道估计不一致"],
        ]
    )
    _add_para(doc, "关键代码：_mobility.py 中的 generate_static / generate_linear / generate_random_walk / generate_random_waypoint 函数。")
    doc.add_page_break()

    # ===== 第 4 章 =====
    _add_heading(doc, "第 4 章  大尺度衰落 —— 链路预算的基础")
    _add_para(doc, (
        "大尺度衰落决定了信号从基站到 UE 的平均功率损失，是整个链路预算的基础。"
        "它直接决定了 SNR（影响信道估计精度）和各小区的接收功率排序（影响服务小区选择和干扰集合）。"
        "同时，LOS/NLOS 的判定结果会传递给第 5 章的 TDL 模型，决定使用 Rician (LOS) 还是 Rayleigh (NLOS) 衰落。"
    ))

    _add_heading(doc, "4.1  路径损耗模型", level=2)
    _add_para(doc, (
        "系统支持三种 3GPP 38.901 规范的场景：\n\n"
        "UMa NLOS（城市宏站）：\n"
        "  PL = 13.54 + 39.08·log10(d_3d) + 20·log10(fc_GHz) - 0.6·(h_UT - 1.5)\n"
        "  阴影衰落标准差 σ_SF = 6 dB\n\n"
        "UMi NLOS（城市微站）：\n"
        "  PL = 22.4 + 35.3·log10(d_3d) + 21.3·log10(fc_GHz) - 0.3·(h_UT - 1.5)\n"
        "  σ_SF = 7.82 dB\n\n"
        "InF（工厂/室内）：\n"
        "  PL = 31.84 + 21.5·log10(d_3d) + 19·log10(fc_GHz)\n"
        "  σ_SF = 4.0 dB"
    ))

    _add_heading(doc, "4.2  LOS/NLOS 概率判定", level=2)
    _add_para(doc, (
        "基于 3GPP 38.901 表 7.4.2-1，根据 2D 距离判定 LOS 概率：\n\n"
        "UMa: P_LOS = min(18/d_2d, 1)·(1 - exp(-d_2d/63)) + exp(-d_2d/63)\n"
        "UMi: P_LOS = min(18/d_2d, 1)·(1 - exp(-d_2d/36)) + exp(-d_2d/36)\n"
        "InF: P_LOS = exp(-d_2d/10)\n\n"
        "判定结果的传递：LOS → TDL-D (Rician, 有直射径 K 因子)；NLOS → TDL-C (Rayleigh, 纯散射)。"
    ))

    _add_heading(doc, "4.3  影响链", level=2)
    _add_table(doc,
        ["输入 (来自上游)", "本环节处理", "输出 (送往下游)"],
        [
            ["d_3d (组网, 第2章)", "路径损耗公式", "PL (dB) → 接收功率"],
            ["scenario (组网, 第2章)", "选择路损公式 + σ_SF", "影响 SNR 分布"],
            ["d_2d (组网, 第2章)", "LOS/NLOS 概率判定", "TDL profile 选择 (第5章)"],
            ["接收功率", "argmax 选择服务小区", "干扰小区集合 (第7章)"],
            ["PL → SNR", "SNR = Tx功率 - PL - 噪声底", "信道估计精度 (第8章)"],
        ]
    )
    doc.add_page_break()

    # ===== 第 5 章 =====
    _add_heading(doc, "第 5 章  小尺度衰落 (TDL) —— 信道矩阵的核心生成")
    _add_para(doc, (
        "TDL (Tapped Delay Line) 信道模型是信道矩阵 h_true 的实际生成器。"
        "它不是独立工作的——它的每一个关键参数都来自前面的环节：天线数来自组网配置，"
        "多普勒频移来自移动性模型，LOS/NLOS 判定来自大尺度衰落，功率缩放来自路径损耗。"
        "TDL 的输出 h_all [K, T, RB, BS_ant, UE_ant] 是后续导频插入、干扰叠加、信道估计的基础数据。"
    ))

    _add_heading(doc, "5.1  多径抽头构建 (PDP)", level=2)
    _add_para(doc, (
        "根据 3GPP 38.901 表 7.5-4/7.5-6 选择 TDL profile (TDL-A~E)。"
        "每个 profile 定义了多径抽头的延迟 τ_l 和归一化功率 P_l。"
        "例如 TDL-C (NLOS 典型) 有 23 个抽头，RMS 延迟扩展约 363ns。"
        "\n\n实际延迟通过乘以场景相关的 τ_rms 缩放：τ_actual = τ_normalized × τ_rms。"
        "功率归一化后 Σ P_l_linear = 1。"
    ))
    _add_image(doc, fig_paths["fig4"], 5.5)

    _add_heading(doc, "5.2  天线空间相关性", level=2)
    _add_para(doc, (
        "采用指数型 ULA (均匀线阵) 模型：R[i,j] = ρ^|i-j|，其中 ρ 为相邻天线相关系数 "
        "(UMa NLOS: 0.7, UMi: 0.5)。\n\n"
        "通过 Cholesky 分解 R = L·L^H 得到着色矩阵 L，将 i.i.d. 高斯信道着色为空间相关信道：\n"
        "  h_corr[l] = L_rx · h_iid[l] · L_tx^T\n\n"
        "天线数 (来自第 2 章组网配置) 决定了相关矩阵的维度和秩，天线数越多，空间选择性越强。"
    ))

    _add_heading(doc, "5.3  Rician LOS 分量", level=2)
    _add_para(doc, (
        "当第 4 章判定为 LOS 时，TDL-D profile 生效，需要叠加直射径：\n"
        "  h = √(K/(1+K)) · h_LOS + √(1/(1+K)) · h_NLOS\n\n"
        "LOS 分量通过导向向量构建：\n"
        "  a_tx[n] = exp(j·2π·0.5·n·sin(θ_aod)) / √N_tx\n"
        "  a_rx[m] = exp(j·2π·0.5·m·sin(θ_aoa)) / √N_rx\n"
        "  h_LOS = a_rx ⊗ a_tx^H"
    ))

    _add_heading(doc, "5.4  多普勒效应", level=2)
    _add_para(doc, (
        "采用改进的 Jakes 和弦和模型，每个多径抽头用 16 个正弦波叠加：\n"
        "  h[l, t] *= Σ_n exp(j·(2π·fd·cos(α_n)·t + φ_n))\n\n"
        "fd 由第 3 章的移动性模型计算得到。fd 越大，信道时变越快，"
        "对后续的信道估计（第 8 章）意味着时域插值的误差越大。"
    ))

    _add_heading(doc, "5.5  时延→频率变换", level=2)
    _add_para(doc, (
        "将时域抽头转换为频域信道响应（每个 RB 一个采样点）：\n"
        "  H[k, t, tx, rx] = Σ_l h_tap[l, t, tx, rx] · exp(-j·2π·k·Δf·τ_l)\n"
        "其中 Δf = 12·SCS（每 RB 12 个子载波），τ_l 为物理延迟（秒）。\n\n"
        "输出维度：h_all [K, T, RB, BS_ant, UE_ant] complex64，K 为小区数。"
    ))

    _add_heading(doc, "5.6  TDL 的上游输入汇总", level=2)
    _add_table(doc,
        ["来自哪个环节", "输入参数", "在 TDL 中的作用"],
        [
            ["组网 (第2章)", "天线数 N_tx, N_rx", "相关矩阵维度和秩"],
            ["移动性 (第3章)", "多普勒 fd", "Jakes 模型的时变速率"],
            ["大尺度 (第4章)", "LOS/NLOS 判定", "选 TDL-D (Rician) 或 TDL-C (Rayleigh)"],
            ["大尺度 (第4章)", "路损→功率缩放", "各小区信道的相对幅度"],
            ["配置", "scenario → τ_rms", "多径时延分布的展宽"],
            ["配置", "bandwidth/SCS → RB 数", "频域采样点数"],
        ]
    )
    doc.add_page_break()

    # ===== 第 6 章 =====
    _add_heading(doc, '第 6 章  TDD 与导频 —— 决定「在哪观测信道」')
    _add_para(doc, (
        "前几章生成了「真实的」信道矩阵 h_true，但实际系统无法直接观测到它。"
        "系统通过在特定时频位置插入已知的导频信号来「探测」信道。"
        "TDD 配置决定了哪些时隙/符号可以放导频（UL 还是 DL），"
        "导频的频率跳频设计决定了能观测到信道的频率范围，"
        "导频密度和梳齿间距决定了观测点的疏密——这些都直接约束了第 8 章信道估计的精度上限。"
    ))

    _add_heading(doc, "6.1  TDD 时隙配置", level=2)
    _add_para(doc, (
        "系统支持 5 种标准 TDD pattern，每种定义了一个周期内时隙的 DL/UL/Special 分配：\n\n"
        "  · DDDSU (5ms)：3 个 DL + 1 Special (10D+2G+2U) + 1 UL\n"
        "  · DDSUU (5ms)：2 DL + 1 Special + 2 UL\n"
        "  · DDDDDDDSUU (10ms)：7 DL + 1 Special + 2 UL\n"
        "  · DSUUD (5ms)：1 DL + 1 Special + 2 UL + 1 DL\n\n"
        "Special 时隙内部进一步分为 DL 符号 (10) + Guard (2) + UL 符号 (2)，共 14 个 OFDM 符号。"
    ))
    _add_para(doc, (
        "TDD pattern 对导频的直接约束：\n"
        "  · SRS (上行导频) 只能放在 UL 符号中 → DDDSU 每周期仅 1 个 UL 时隙 + Special 中 2 个 UL 符号\n"
        "  · CSI-RS (下行导频) 只能放在 DL 符号中\n"
        "  · BOTH (双向) 模式下，UL/DL 导频分别放在各自方向的符号中"
    ))
    _add_image(doc, fig_paths["fig5"], 6.0)

    _add_heading(doc, "6.2  SRS 上行导频", level=2)
    _add_para(doc, (
        "SRS (Sounding Reference Signal) 是上行信道探测的核心，遵循 3GPP TS 38.211 §6.4.1.4。"
    ))

    _add_heading(doc, "6.2.1  ZC 基础序列", level=3)
    _add_para(doc, (
        "SRS 使用 Zadoff-Chu (ZC) 序列作为基础：\n"
        "  x_u(n) = exp(-j·π·u·n·(n+1) / N_ZC),  n = 0..N_ZC-1\n\n"
        "其中 N_ZC 为不超过 M_sc 的最大素数，u 为根序号 (由组号和序列号确定)。"
        "ZC 序列的关键特性：恒模、零自相关、低互相关——这使得不同小区/UE 可用不同根号实现部分正交。\n\n"
        "短序列 (M_sc ∈ {6,12,18,24}) 使用查表方式：30 组 × M_sc 列的相位表 φ(n)，"
        "序列为 exp(j·π·φ(n)/4)。"
    ))

    _add_heading(doc, "6.2.2  频率跳频与带宽树", level=3)
    _add_para(doc, (
        "SRS 频率跳频是系统设计中的关键机制——它决定了信道估计能「看到」多大的频率范围。\n\n"
        "带宽树机制 (3GPP Table 6.4.1.4.3-1)：\n"
        "  · C_SRS 参数选择一行配置，定义 4 级带宽树 m_SRS[0..3] 和分频因子 N[0..3]\n"
        "  · Level 0 为最大带宽 (覆盖全频带 ~273 RB)，逐级细分\n"
        "  · 每次发送 SRS 占用 m_SRS[B_hop] 个 RB (一个子频带)\n"
        "  · 跳频索引 n_SRS 随时隙递增，驱动 F_b 在各子频带间循环\n\n"
        "累积 RB 覆盖：经过一个完整跳频周期 (N[B_hop] 次)，SRS 累积覆盖了上一级的全部 RB，"
        "从而实现全频带信道探测——但代价是需要多个时隙才能完成。"
    ))
    _add_para(doc, (
        "跳频与移动性的耦合 (回顾第 3 章)：\n"
        "  · 信道相干时间 Tc ≈ 0.423/fd\n"
        "  · 跳频完成全频带覆盖需要 N_hop 个 SRS 周期 × T_SRS 间隔\n"
        "  · 若 N_hop × T_SRS > Tc，不同频段的估计对应不同时刻的信道 → 时间不一致\n"
        "  · 因此高速 UE (大 fd) 场景下需要更短的 SRS 周期或更少的跳频级数"
    ))
    _add_image(doc, fig_paths["fig6"], 5.8)

    _add_heading(doc, "6.2.3  梳齿映射", level=3)
    _add_para(doc, (
        "SRS 在频域按梳齿 (comb) 方式映射到子载波，K_TC ∈ {2, 4, 8}：\n"
        "  · K_TC = 2：每隔 1 个子载波放一个 SRS 符号\n"
        "  · K_TC = 4/8：更稀疏\n\n"
        "K_TC 越大 → 单次 SRS 占用更多 RB 但子载波密度更低 → 频域分辨率降低 → 影响信道估计的频率精度。"
        "不同 UE 可用不同梳齿偏移实现 CDM 复用。"
    ))

    _add_heading(doc, "6.3  CSI-RS 下行导频", level=2)
    _add_para(doc, (
        "CSI-RS 遵循 3GPP TS 38.211 §7.4.1.5，使用 Gold 序列生成 QPSK 符号：\n"
        "  c_init = (2^10·(N_symb·n_s + l + 1)·(2·N_ID + 1) + N_ID) mod 2^31\n"
        "  r(m) = (1/√2)·[1-2c(2m)] + j·(1/√2)·[1-2c(2m+1)]\n\n"
        "支持多种 CDM 模式：noCDM (1 RE/端口)、fd-CDM2 (2 频域 RE)、cdm4 (2×2 频/时)、"
        "cdm8 (2×4 频/时)。密度选项：ρ = 0.5 / 1.0 / 3.0。\n\n"
        "N_ID 由小区 PCI 决定 (来自第 2 章组网)，确保邻区 CSI-RS 序列不同——但不完全正交，"
        "这正是第 7 章 DL 干扰的来源。"
    ))

    _add_heading(doc, "6.4  SSB 同步信号块", level=2)
    _add_para(doc, (
        "SSB 用于小区搜索和初始同步，包含 4 个 OFDM 符号 × 240 个子载波：\n"
        "  · Symbol 0: PSS (主同步信号, 127 子载波)\n"
        "  · Symbol 1: PBCH-DMRS + PBCH 数据\n"
        "  · Symbol 2: SSS (辅同步信号, 127 子载波)\n"
        "  · Symbol 3: PBCH-DMRS + 数据\n\n"
        "通过 DFT 波束扫描在多个方向发送 SSB，UE 测量各波束/各小区的 SS-RSRP，"
        "选择最强波束。这一测量在第 9 章详述。"
    ))
    doc.add_page_break()

    # ===== 第 7 章 =====
    _add_heading(doc, "第 7 章  干扰建模 —— 多小区环境的核心挑战")
    _add_para(doc, (
        "在单小区仿真中，信道估计只需对抗热噪声。但在多小区环境下，"
        "邻区的导频信号（SRS 或 CSI-RS）通过各自的信道到达本小区接收端，"
        "与服务小区的导频叠加形成干扰。干扰的强度由组网 (第 2 章确定的邻区集合) 和 "
        "大尺度衰落 (第 4 章确定的各邻区路径损耗) 共同决定；"
        "干扰的形态由导频设计 (第 6 章的不同 ZC 根/Gold 种子) 决定。"
    ))

    _add_heading(doc, "7.1  UL 干扰模型 (基站接收端)", level=2)
    _add_para(doc, (
        "在上行链路中，基站接收的信号为：\n"
        "  Y = H_serving · X_SRS_serving + Σ_k H_intf_k · X_SRS_k + noise\n\n"
        "其中 X_SRS_k 是第 k 个邻区 UE 的 SRS 导频。不同 UE 使用不同的 ZC 根号：\n"
        "  u_k = ((cell_id_k + ue_index + 1)·7 + 3) mod N_ZC\n\n"
        "由于 ZC 序列在不同根号间不完全正交（互相关不为零），邻区 SRS 构成非高斯干扰。"
        "每个邻区随机激活 0 到 num_interfering_ues 个干扰 UE，模拟实际网络的负载随机性。"
    ))

    _add_heading(doc, "7.2  DL 干扰模型 (UE 接收端)", level=2)
    _add_para(doc, (
        "在下行链路中，UE 接收：\n"
        "  Y = H_serving · X_CSIRS_serving + Σ_k H_intf_k · X_CSIRS_k + noise\n\n"
        "邻区 CSI-RS 由各自的 cell_id 作为种子生成不同的 Gold 序列。"
        "随机选取邻区子集激活 CSI-RS 发送。"
    ))

    _add_heading(doc, "7.3  SIR/SINR 计算", level=2)
    _add_para(doc, (
        "  P_serving = mean(|H_serving · X_serving|²)\n"
        "  P_intf = mean(|Σ H_intf_k · X_intf_k|²)\n"
        "  SIR = 10·log10(P_serving / P_intf)\n"
        "  SINR = 10·log10(P_serving / (P_intf + P_noise))\n\n"
        "SIR 和 SINR 作为元数据写入 ChannelSample，同时直接影响信道估计质量：\n"
        "  · 低 SIR → LS 估计的偏差大 (干扰被当作噪声)\n"
        "  · MMSE 通过先验协方差可以部分抑制干扰，但在极低 SIR 下也会退化"
    ))

    _add_heading(doc, "7.4  干扰的因果链", level=2)
    _add_table(doc,
        ["上游环节", "提供的输入", "对干扰的影响"],
        [
            ["组网 (第2章)", "邻区集合 (K-1 个小区)", "干扰源数量"],
            ["大尺度 (第4章)", "各邻区的路径损耗", "干扰信号到达功率"],
            ["TDL (第5章)", "邻区信道矩阵 h_intf", "干扰的空间/频率选择性"],
            ["导频 (第6章)", "邻区导频序列 X_intf", "干扰的形态（非正交残余）"],
            ["配置", "num_interfering_ues", "每邻区激活的干扰 UE 数"],
        ]
    )
    _add_para(doc, "关键代码：_interference_estimation.py 中的 estimate_channel_with_interference() 和 estimate_paired_channels()。")
    doc.add_page_break()

    # ===== 第 8 章 =====
    _add_heading(doc, "第 8 章  信道估计 —— 从观测恢复信道")
    _add_para(doc, (
        "信道估计是整条因果链的汇聚点——前面所有环节的输出都在这里产生影响。"
        "SNR (来自路径损耗) 决定噪声水平，SIR (来自干扰) 决定偏差大小，"
        "导频位置 (来自 TDD 和 SRS/CSI-RS 设计) 决定观测点的时频分布，"
        "TDL 的延迟扩展 τ_rms 决定 MMSE 先验协方差是否准确。"
    ))

    _add_heading(doc, "8.1  三种估计模式", level=2)
    _add_para(doc, (
        "· ideal：跳过所有估计步骤，直接使用 h_true。这是性能上界基准。\n\n"
        "· ls_linear：LS 最小二乘 + 线性插值。\n"
        "  Ĥ_LS[k] = Y_RS[k] · X_RS*[k] / |X_RS[k]|²\n"
        "  无偏估计，但方差为 σ²/|X|²，在低 SNR 下噪声放大严重。\n\n"
        "· ls_mmse：LS 估计 + MMSE 细化 + 线性插值。\n"
        "  Ĥ_MMSE = R_hh · (R_hh + (1/SNR)·I)^-1 · Ĥ_LS\n"
        "  利用信道的频率相关性（指数 PDP 先验）来抑制噪声，在低 SNR 下增益显著。"
    ))

    _add_heading(doc, "8.2  MMSE 先验协方差", level=2)
    _add_para(doc, (
        "MMSE 需要信道频域相关矩阵 R_hh 作为先验。系统采用指数 PDP 模型：\n"
        "  R[k, l] = 1 / (1 + j·2π·(k-l)·Δf·τ_rms)\n\n"
        "其中 Δf 为导频间的频率间隔，τ_rms 为 RMS 延迟扩展 (来自第 5 章 TDL 的配置)。"
        "如果 τ_rms 设置不准确，协方差矩阵与实际信道不匹配，MMSE 性能会退化。"
    ))

    _add_heading(doc, "8.3  2D 分离插值", level=2)
    _add_para(doc, (
        "LS/MMSE 只在导频位置得到估计值，需要插值到全部 RB × 全部 OFDM 符号：\n"
        "  1. 频率方向：在同一时隙内，将导频 RB 位置的估计值线性插值到所有 RB\n"
        "  2. 时间方向：在同一频率上，将不同时隙的估计值线性插值到所有 OFDM 符号\n\n"
        "插值误差的两个主要来源：\n"
        "  · 频域：导频密度不够 (K_TC 大 或 SRS 跳频未覆盖) → 频率选择性信道的细节丢失\n"
        "  · 时域：UE 移动快 (fd 大) → 相邻时隙间信道变化大 → 线性内插假设不准"
    ))
    _add_image(doc, fig_paths["fig7"], 5.8)

    _add_heading(doc, "8.4  估计精度的决定因素汇总", level=2)
    _add_table(doc,
        ["影响因素", "来源", "对估计精度的影响"],
        [
            ["SNR", "路损 (第4章) + 发射功率", "SNR↓ → LS 噪声↑ → MMSE 增益↑"],
            ["SIR", "干扰 (第7章)", "SIR↓ → 干扰当噪声 → 估计偏差"],
            ["导频密度", "SRS K_TC / CSI-RS density (第6章)", "密度↓ → 观测点少 → 插值误差↑"],
            ["跳频覆盖", "SRS 跳频 (第6章)", "未覆盖频段无直接观测 → 纯插值"],
            ["信道时变", "移动性 (第3章) → fd", "fd↑ → 时域插值误差↑"],
            ["MMSE 先验", "TDL τ_rms (第5章)", "τ_rms 不匹配 → 协方差失准"],
        ]
    )
    _add_para(doc, "关键文件：channel_est/pipeline.py (统一入口), ls.py, mmse.py, interpolate.py。")
    doc.add_page_break()

    # ===== 第 9 章 =====
    _add_heading(doc, "第 9 章  SSB 测量与多小区指标")
    _add_para(doc, (
        "SSB 测量独立于 SRS/CSI-RS 的信道估计链路，但同样依赖组网 (多小区信道) 和天线配置 (波束数)。"
    ))

    _add_heading(doc, "9.1  波束扫描", level=2)
    _add_para(doc, (
        "基站通过 DFT 波束在多个方向发送 SSB：\n"
        "  beam[a, b] = exp(j·2π·0.5·sin(θ_b)·a) / √N_ant\n"
        "  θ_b = π·(b/num_beams - 0.5)\n\n"
        "波束数与天线数相关 (来自第 2 章)，天线越多 → 波束越窄 → 空间分辨率越高。"
    ))

    _add_heading(doc, "9.2  测量指标 (3GPP 38.215)", level=2)
    _add_para(doc, (
        "对每个小区 k 的每个波束 b 计算：\n"
        "  · SS-RSRP (dBm)：SSS 符号上的平均接收功率，选最强波束报告\n"
        "  · RSRQ (dB)：N·RSRP / RSSI，反映信道质量\n"
        "  · SS-SINR (dB)：RSRP_serving / (Σ RSRP_interferer + noise)\n\n"
        "这些指标作为 ChannelSample 的 ssb_rsrp_dBm / ssb_rsrq_dB / ssb_sinr_dB 字段输出，"
        "为下游的 AI 模型提供多小区级别的信号质量信息。"
    ))
    doc.add_page_break()

    # ===== 第 10 章 =====
    _add_heading(doc, "第 10 章  三种采集源的建模差异")
    _add_para(doc, (
        "三种采集源在上述各环节中的实现方式不同，但最终都输出统一的 ChannelSample 契约。"
    ))
    _add_table(doc,
        ["建模环节", "InternalSim (Python)", "Sionna RT (射线追踪)", "QuaDRiGa (MATLAB)"],
        [
            ["组网", "Python hex_grid 六边形", "Sionna Scene 3D场景", "MATLAB gen_hex_positions"],
            ["路径损耗", "3GPP 38.901 公式", "射线追踪物理传播", "QuaDRiGa 内置模型"],
            ["小尺度衰落", "TDL 抽头模型 (NumPy)", "射线追踪 CIR → FFT", "QuaDRiGa 信道生成器"],
            ["导频", "Python srs.py/csi_rs.py", "同左 (共享模块)", "MATLAB ul/dl pipeline"],
            ["干扰", "Python _interference_est", "同左 (共享模块)", "MATLAB apply_interference"],
            ["信道估计", "Python channel_est/", "同左 (共享模块)", "MATLAB dl_csirs_pipeline"],
            ["输出", "ChannelSample (直接)", "ChannelSample (直接)", ".mat → Python 转换 → ChannelSample"],
        ]
    )

    _add_heading(doc, "10.1  TDD 互易性 (paired 模式)", level=2)
    _add_para(doc, (
        "当 link = BOTH 时，系统生成配对的 UL/DL 信道。基于 TDD 互易性：\n"
        "  H_DL = conj(H_UL^T) + 校准误差\n\n"
        "配对模式下，UL 和 DL 分别带各自方向的干扰 (SRS 干扰 vs CSI-RS 干扰)，"
        "对应 ul_sir_dB 和 dl_sir_dB 两个独立的指标。"
    ))
    doc.add_page_break()

    # ===== 第 11 章 =====
    _add_heading(doc, "第 11 章  数据输出 —— ChannelSample 契约")
    _add_para(doc, (
        "ChannelSample 是所有环节的最终汇聚点，也是下游特征提取和 AI 模型训练的唯一输入契约。"
        "每个字段都可以追溯到上游的具体建模环节。"
    ))

    _add_table(doc,
        ["字段", "维度", "物理含义", "来源环节"],
        [
            ["h_serving_true", "[T, RB, BS_ant, UE_ant]", "理想服务小区信道", "TDL (第5章)"],
            ["h_serving_est", "[T, RB, BS_ant, UE_ant]", "估计的服务小区信道", "信道估计 (第8章)"],
            ["h_interferers", "[K-1, T, RB, BS_ant, UE_ant]", "干扰小区信道 (可选)", "TDL + 干扰 (第5+7章)"],
            ["snr_dB / sir_dB / sinr_dB", "标量", "链路级质量指标", "大尺度 (第4章) + 干扰 (第7章)"],
            ["ssb_rsrp_dBm", "[K]", "各小区最佳波束 SS-RSRP", "SSB 测量 (第9章)"],
            ["ssb_rsrq_dB / ssb_sinr_dB", "[K]", "RSRQ / SS-SINR", "SSB 测量 (第9章)"],
            ["link", "'UL' | 'DL'", "链路方向", "TDD (第6章)"],
            ["channel_est_mode", "'ideal'|'ls_linear'|'ls_mmse'", "估计模式", "信道估计 (第8章)"],
            ["tdd_pattern", "str 'DDDSU'等", "TDD 配置", "TDD (第6章)"],
            ["ue_position", "[3]", "UE 三维位置 (m)", "组网 (第2章) + 移动性 (第3章)"],
            ["channel_model", "str 'TDL-C'等", "信道模型名称", "大尺度 (第4章) → TDL 选择"],
            ["meta", "dict", "完整配置快照", "全部环节"],
        ]
    )

    _add_para(doc, (
        "序列化：每个 ChannelSample 保存为 .pt 文件 (PyTorch 格式)，"
        "标量元数据同步写入 manifest.parquet 索引文件。"
        "复数数组在序列化时拆分为 (real, imag) float32 对以兼容 JSON。"
    ))

    return doc


# ===========================================================================
# Main
# ===========================================================================
def main():
    print("正在生成插图...")
    fig_paths = {
        "fig1": fig1_causal_chain(),
        "fig2": fig2_hex_topology(),
        "fig3": fig3_mobility(),
        "fig4": fig4_tdl_pdp(),
        "fig5": fig5_tdd_pattern(),
        "fig6": fig6_srs_hopping(),
        "fig7": fig7_channel_est(),
    }
    print(f"  7 张插图已生成到 {IMG_DIR}")

    print("正在组装 Word 文档...")
    doc = build_document(fig_paths)
    doc.save(str(DOCX_PATH))
    print(f"  文档已保存: {DOCX_PATH}")

    # 清理临时图片
    for p in IMG_DIR.glob("*.png"):
        p.unlink()
    IMG_DIR.rmdir()
    print("  临时文件已清理。完成！")


if __name__ == "__main__":
    main()
